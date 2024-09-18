Створення файлів для тестування.py
import os
import random
import zipfile
from PIL import Image
import pandas as pd
from fpdf import FPDF
import docx

# Шлях до директорії всередині вашого проекту
output_dir = os.path.join('C:\\Путь\\Путь\\Путь\\Путь', 'test_files')
os.makedirs(output_dir, exist_ok=True)

# Функція для створення випадкового тексту
def random_text(size=1024):
    return ''.join(random.choices('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789', k=size))

# 1. Створення файлу .txt
with open(os.path.join(output_dir, 'test.txt'), 'w') as f:
    f.write(random_text(1024 * 1024))  # Запис 1 МБ випадкового тексту

# 2. Створення .doc файлу (старий формат MS Word)
doc_content = random_text(512)  # Контент-заповнювач
with open(os.path.join(output_dir, 'test.doc'), 'w') as f:
    f.write(doc_content * 2000)  # Зробити файл приблизно 1 МБ

doc = docx.Document()
doc.add_paragraph(random_text(512))
for _ in range(2000):
    doc.add_paragraph(random_text(512))
doc.save(os.path.join(output_dir, 'test.docx'))

data = {'Column1': [random_text(512) for _ in range(2000)]}
df = pd.DataFrame(data)
df.to_excel(os.path.join(output_dir, 'test.xlsx'), index=False)

df.to_csv(os.path.join(output_dir, 'test.csv'), index=False)

xml_content = f"<root>{random_text(1024 * 1024)}</root>"
with open(os.path.join(output_dir, 'test.xml'), 'w') as f:
    f.write(xml_content)

pdf = FPDF()
pdf.add_page()
pdf.set_font('Arial', size=12)
for _ in range(4000):
    pdf.cell(200, 10, txt=random_text(50), ln=True)
pdf.output(os.path.join(output_dir, 'test.pdf'))

image_size = (1024, 1024)
colors = [random.randint(0, 255) for _ in range(3)]
image = Image.new('RGB', image_size, tuple(colors))
image.save(os.path.join(output_dir, 'test.jpg'))
image.save(os.path.join(output_dir, 'test.jpeg'))
image.save(os.path.join(output_dir, 'test.png'))
image.save(os.path.join(output_dir, 'test.gif'))

with zipfile.ZipFile(os.path.join(output_dir, 'test.zip'), 'w') as zipf:
    zipf.write(os.path.join(output_dir, 'test.txt'), arcname='test.txt')

print(os.listdir(output_dir))
